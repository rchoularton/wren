"""
Markdown-to-Word document builder.

Converts markdown content into a properly formatted Word document
using journal-specific styling.
"""

import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styles import apply_journal_styles, JOURNAL_STYLES


class MarkdownSection:
    """Represents a section of markdown content."""

    def __init__(self, level: int, title: str, content: str):
        self.level = level  # 0=title, 1=h1, 2=h2, etc.
        self.title = title
        self.content = content


def parse_markdown(content: str) -> list[MarkdownSection]:
    """Parse markdown into sections."""
    sections = []
    lines = content.split('\n')

    current_section = None
    current_content = []

    for line in lines:
        # Check for headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)

        if heading_match:
            # Save previous section
            if current_section is not None:
                current_section.content = '\n'.join(current_content).strip()
                sections.append(current_section)
                current_content = []

            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            current_section = MarkdownSection(level, title, '')
        else:
            current_content.append(line)

    # Save final section
    if current_section is not None:
        current_section.content = '\n'.join(current_content).strip()
        sections.append(current_section)
    elif current_content:
        # Content before any heading (title/metadata)
        sections.insert(0, MarkdownSection(0, '', '\n'.join(current_content).strip()))

    return sections


def split_paragraphs(text: str) -> list[str]:
    """Split text into paragraphs."""
    # Split on double newlines, filter empty
    paragraphs = re.split(r'\n\s*\n', text)
    return [p.strip() for p in paragraphs if p.strip()]


def add_formatted_paragraph(doc: Document, text: str, style_name: str):
    """Add a paragraph with inline formatting (bold, italic, references)."""
    p = doc.add_paragraph(style=style_name)

    # Handle inline formatting
    # Pattern: **bold**, *italic*, [number] references
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|\[\d+(?:,\s*\d+)*\])'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            run = p.add_run(part[1:-1])
            run.italic = True
        elif re.match(r'\[\d+(?:,\s*\d+)*\]', part):
            # Reference citations - keep as-is
            p.add_run(part)
        else:
            p.add_run(part)

    return p


def build_word_document(markdown_path: Path, output_path: Path, journal: str) -> Path:
    """
    Build a Word document from markdown.

    Args:
        markdown_path: Path to source markdown file
        output_path: Path for output Word document
        journal: Journal style to apply (e.g., 'default', 'nature_food')

    Returns:
        Path to generated Word document
    """
    if journal not in JOURNAL_STYLES:
        raise ValueError(f"Unknown journal: {journal}")

    # Read markdown
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Parse into sections
    sections = parse_markdown(content)

    # Create document with journal styles
    doc = Document()
    apply_journal_styles(doc, journal)

    # Track section context for styling
    in_methods = False
    in_references = False
    in_abstract = False

    for section in sections:
        # Handle pre-heading content (title block)
        if section.level == 0 and section.content:
            # This is the title/author block before any heading
            paragraphs = split_paragraphs(section.content)
            for i, para in enumerate(paragraphs):
                if i == 0:
                    # First paragraph is title
                    p = doc.add_paragraph(para, style='Title')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Author/affiliation info
                    p = doc.add_paragraph(para, style='Author')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
            continue

        # Detect special sections
        title_lower = section.title.lower()
        if 'abstract' in title_lower:
            in_abstract = True
            in_methods = False
            in_references = False
        elif 'method' in title_lower:
            in_abstract = False
            in_methods = True
            in_references = False
        elif 'reference' in title_lower:
            in_abstract = False
            in_methods = False
            in_references = True
        elif section.level == 1:
            in_abstract = False
            # Keep methods flag if we're in a subsection

        # Add heading
        if section.title:
            if section.level == 1:
                # Add page break before major sections
                if title_lower in ['results', 'methods', 'references', 'discussion']:
                    doc.add_page_break()

                p = doc.add_paragraph(section.title, style='Heading1')
            elif section.level >= 2:
                p = doc.add_paragraph(section.title, style='Heading2')

        # Add content paragraphs
        if section.content:
            paragraphs = split_paragraphs(section.content)

            for i, para in enumerate(paragraphs):
                if in_abstract:
                    add_formatted_paragraph(doc, para, 'Abstract')
                elif in_references:
                    add_formatted_paragraph(doc, para, 'Reference')
                elif i == 0:
                    # First paragraph in section - no indent
                    add_formatted_paragraph(doc, para, 'BodyTextNoIndent')
                else:
                    # Subsequent paragraphs - with indent
                    add_formatted_paragraph(doc, para, 'BodyText')

    # Save document
    doc.save(output_path)
    return output_path


def build_paper(paper_id: str, version: int = None) -> Path:
    """
    Build Word document for a paper using its config.

    Args:
        paper_id: Paper identifier (e.g., 'paper1a')
        version: Optional version number for output filename

    Returns:
        Path to generated Word document
    """
    from .paths import get_paper_config, get_drafts_dir, get_latest_markdown

    config = get_paper_config(paper_id)
    drafts_dir = get_drafts_dir(paper_id)
    journal = config.get('journal', 'default')

    # Find markdown source
    md_path = get_latest_markdown(paper_id)
    if not md_path:
        raise FileNotFoundError(f"No markdown file found in {drafts_dir}")

    # Determine output filename
    if version is None:
        version = config.get('current_version', 1)

    output_name = f"{config.get('title', paper_id).replace(' ', '_')}_v{version}.docx"
    output_path = drafts_dir / output_name

    return build_word_document(md_path, output_path, journal)
